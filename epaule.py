import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import docx
from docx.shared import Inches
import os
import docx
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from tabula import read_pdf
from tabulate import tabulate
from statistics import mean
import numpy as np
import matplotlib.font_manager as fm

###################################################################
# CODE POUR TRAITER LES QUESTIONNAIRES SUBJECTIFS
###################################################################

# Obtenir le chemin d'accès absolu du répertoire courant
current_directory = os.path.dirname(os.path.abspath(__file__))

# Construire le chemin d'accès complet vers le fichier de police
font_path = os.path.join(current_directory, "Cambria-Font-For-Windows.ttf")

# Définir les propriétés de la police
font = FontProperties(fname=font_path)

# Fonction pour calculer le pourcentage des réponses
def calculate_percentage(question):
    total_responses = len(question)
    response_counts = question.value_counts()
    response_percentages = (response_counts / total_responses) * 100
    return response_percentages

# Définir les couleurs pour chaque réponse
colors = {
    "Tout à fait d'accord": '#549e39',
    "Plutôt d'accord": '#8ab833',
    "Moyennement d'accord": '#c0cf3a',
    "Plutôt pas d'accord": '#029676',
    "Pas du tout d'accord": '#4ab5c4'
}

# Charger le fichier Excel et lire les données du questionnaire dans un DataFrame
df1 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="F:G", header=None)
df2 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="H:J", header=None)
df3 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="K:M", header=None)
df4 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="N:Q", header=None)

# Supprimer la première ligne (questions) et la remplacer par les colonnes appropriées
df1 = df1.drop(0)
df1.columns = ['Question1', 'Question2']
df2 = df2.drop(0)
df2.columns = ['Question3', 'Question4', 'Question5']
df3 = df3.drop(0)
df3.columns = ['Question6', 'Question7', 'Question8'] 
df4 = df4.drop(0)
df4.columns = ['Question9', 'Question10', 'Question11', 'Question12'] 

# Renommer les réponses contenant des caractères spéciaux
df1.replace({"Plutôt d’accord": "Plutôt d'accord",
             "Pas du tout d’accord": "Pas du tout d'accord",
             "Moyennement d’accord": "Moyennement d'accord",
             "Tout à fait d'accord": "Tout à fait d'accord",
             "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)

# Calculer les pourcentages de réponses pour chaque question
question1_percentages = calculate_percentage(df1['Question1'])
question2_percentages = calculate_percentage(df1['Question2'])
question_percentages_2 = {}
for col in df2.columns:
    df2[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df2.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_2[col] = calculate_percentage(df2[col])

question_percentages_3 = {} # Ajouter cette ligne
for col in df3.columns: # Ajouter cette ligne
    df3[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df3.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_3[col] = calculate_percentage(df3[col]) # Ajouter cette ligne

question_percentages_4 = {} # Ajouter cette ligne
for col in df4.columns: # Ajouter cette ligne
    df4[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df4.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_4[col] = calculate_percentage(df4[col]) # Ajouter cette ligne
    
# Créer les subplots
cm = 1/2.54  # centimeters in inches
fig1, (ax1, ax2) = plt.subplots(2, 1, figsize=(20*cm, 30*cm))
fig2, (ax3, ax4) = plt.subplots(2, 1, figsize=(20*cm, 30*cm))


# Gérer le premier subplot
handles = []
labels = []
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

# Première question
bottom_values_q1 = [0] * len(question1_percentages)
for i, (category, percentage) in enumerate(question1_percentages.items()):
    ax1.bar('Question 1', percentage, bottom=bottom_values_q1, color=colors.get(category, 'lightgray'))
    bottom_values_q1 = [sum(x) for x in zip(bottom_values_q1, [percentage])]

# Deuxième question
bottom_values_q2 = [0] * len(question2_percentages)
for i, (category, percentage) in enumerate(question2_percentages.items()):
    ax1.bar('Question 2', percentage, bottom=bottom_values_q2, color=colors.get(category, 'lightgray'))
    bottom_values_q2 = [sum(x) for x in zip(bottom_values_q2, [percentage])]

#ax1.set_xlabel('Questions', fontproperties=font)
ax1.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax1.set_title('Mise en place et réglages de l\'exosquelette', fontproperties=title_font)

ax1.legend(handles, labels, loc='upper left', bbox_to_anchor=(1.0, 1.0), prop=font)
ax1.grid(True)

ax1.set_xticks([0, 1])
ax1.set_xticklabels(['Facilité de mise en place', 'Adaptation à la morphologie'], fontproperties=font)

# Gérer le deuxième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q3_to_q7 = [0] * len(df2.columns)
for col, percentages in enumerate(question_percentages_2.values()):
    bottom_values_q3_to_q7 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):
        ax2.bar('Question {}'.format(col+3), percentage, bottom=bottom_values_q3_to_q7, color=colors.get(category, 'lightgray'))
        bottom_values_q3_to_q7 = [sum(x) for x in zip(bottom_values_q3_to_q7, [percentage])]

#ax2.set_xlabel('Questions', fontproperties=font)
ax2.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax2.set_title("L'utilisation de l'exosquelette", fontproperties=title_font)

ax2.grid(True)

ax2.set_xticks(range(len(df2.columns)))
ax2.set_xticklabels(['Facilité des mouvements', 'Facilité de déplacement', 'Habituation'], fontproperties=font)

# Enregistrer et afficher la première figure
fig1.tight_layout()
fig1.savefig('subplot1.png')
#plt.show()

# Gérer le troisième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q8_to_q10 = [0] * len(df3.columns)
for col, percentages in enumerate(question_percentages_3.values()):
    bottom_values_q8_to_q10 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):
        ax3.bar('Question {}'.format(col+8), percentage, bottom=bottom_values_q8_to_q10, color=colors.get(category, 'lightgray'))
        bottom_values_q8_to_q10 = [sum(x) for x in zip(bottom_values_q8_to_q10, [percentage])]

#ax3.set_xlabel('Questions', fontproperties=font)
ax3.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax3.set_title("Le confort de l'exosquelette", fontproperties=title_font)

ax3.grid(True)

ax3.set_xticks(range(len(df3.columns)))
ax3.set_xticklabels(['Confort', 'Frottements', 'Me donne chaud'], fontproperties=font)

# Gérer le quatrième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q11_to_q14 = [0] * len(df4.columns)
for col, percentages in enumerate(question_percentages_4.values()):
    bottom_values_q11_to_q14 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):
        ax4.bar('Question {}'.format(col+11), percentage, bottom=bottom_values_q11_to_q14, color=colors.get(category, 'lightgray'))
        bottom_values_q11_to_q14 = [sum(x) for x in zip(bottom_values_q11_to_q14, [percentage])]

#ax4.set_xlabel('Questions', fontproperties=font)
ax4.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax4.set_title("Le ressenti avec l'exosquelette", fontproperties=title_font)

ax4.grid(True)

ax4.set_xticks(range(len(df4.columns)))
ax4.set_xticklabels(['Je me sens nerveux', 'J\'aime travailler avec', 'Je me sens isolé avec','Je me sens confiant avec'], fontproperties=font)

# Enregistrer et afficher la deuxième figure
fig2.tight_layout()
fig2.savefig('subplot2.png')
#plt.show()


###################################################################
# CODE POUR IMPORTER LA PHOTO
###################################################################
def replace_photo_tag_with_image(doc_path, image_path, new_doc_path):
    # Ouvrir le document Word
    doc = docx.Document(doc_path)

    # Parcourir les paragraphes du document
    for paragraph in doc.paragraphs:
        # Vérifier si le paragraphe contient la balise ##photo##
        if "##photo##" in paragraph.text:
            # Remplacer la balise par une image
            run = paragraph.runs[0]
            run.text = ""  # Effacer la balise
            run.add_picture(image_path, width=Inches(5))  # Ajouter l'image

    # Enregistrer les modifications dans le nouveau document
    doc.save(new_doc_path)

# Créer une fenêtre tkinter cachée pour ouvrir la boîte de dialogue
root = tk.Tk()
root.withdraw()
root.deiconify()  # Forcer l'affichage de la fenêtre

# Ouvrir la boîte de dialogue pour sélectionner l'image
image_path = filedialog.askopenfilename(title="Sélectionner une image", filetypes=[("Fichiers image", "*.png *.jpg *.jpeg")])

# Vérifier si une image a été sélectionnée
if image_path:
    # Appeler la fonction pour remplacer la balise ##photo## par l'image dans un nouveau document
    new_doc_path = "nouveau_document.docx"
    replace_photo_tag_with_image("CR_epaules2.docx", image_path, new_doc_path)
    print(f"L'image a été insérée avec succès dans le nouveau document : {new_doc_path}")
else:
    print("Aucune image sélectionnée.")

# Fermer la fenêtre tkinter cachée
root.destroy()

###################################################################
# REMPLISSAGE DU RAPPORT
###################################################################

def remplacer_balises(document, balises):
    """
    Remplace les balises dans le document par les valeurs correspondantes.

    Arguments :
    document -- Document Word (obj)
    balises -- Dictionnaire contenant les balises en tant que clés et leurs valeurs correspondantes en tant que valeurs (dict)
    """
    for paragraph in document.paragraphs:
        for balise, valeur in balises.items():
            if balise in paragraph.text:
                paragraph.text = paragraph.text.replace(balise, valeur)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for balise, valeur in balises.items():
                        if balise in paragraph.text:
                            paragraph.text = paragraph.text.replace(balise, valeur)
                            
                            
# Obtenez le chemin absolu du dossier contenant les fichiers PDF
#directory = r'C:\Users\l.hoffner\OneDrive\Documents\ErgoSante\Documents_compte_rendu\AVEC'

def choisir_dossier(dossier_initial):
    """
    Ouvre une fenêtre de dialogue pour sélectionner un dossier et renvoie le chemin du dossier sélectionné.

    Arguments :
    dossier_initial -- Chemin du dossier de départ (str)
    """
    root = tk.Tk()
    root.withdraw()  # Cache la fenêtre principale de Tkinter
    directory = filedialog.askdirectory(initialdir=dossier_initial, title="Sélectionnez le dossier contenant les fichiers PDF")
    if directory:
        return directory
    else:
        messagebox.showwarning("Avertissement", "Aucun dossier sélectionné. L'application va se fermer.")
        root.destroy()
        exit()

# Ouvrir la boîte de dialogue pour sélectionner un dossier
directory = filedialog.askdirectory(title="Sélectionner un dossier")

# Vérifier si un dossier a été sélectionné
if directory:
    # Obtenez le chemin absolu du dossier sélectionné
    pdf_folder = os.path.abspath(directory)
    print(f"Chemin absolu du dossier sélectionné : {pdf_folder}")
else:
    print("Aucun dossier sélectionné.")

#pdf_folder = os.path.abspath(directory)

seventh_tables_data = []
eighth_tables_data = []
ninth_tables_data = []
tenth_tables_data = []
fifteenth_tables_data = []
sixteenth_tables_data = []



# Parcourez tous les fichiers PDF dans le dossier et lisez-les
for file_name in os.listdir(pdf_folder):
    if file_name.endswith('.pdf'):
        file_path = os.path.join(pdf_folder, file_name)
        # Lecture des tables à partir du fichier PDF
        dfs = read_pdf(file_path, pages='all', multiple_tables=True)
        # Ajouter le premier tableau de chaque fichier à la liste
        seventh_table = dfs[6].iloc[:, 1]
        eighth_table = dfs[7].iloc[:, 1]
        ninth_table = dfs[8].iloc[:, 1]
        tenth_table = dfs[9].iloc[:, 1]

        fifteenth_table = dfs[14].iloc[:, 1]
        sixteenth_table = dfs[15].iloc[:, 1]
        
        
        liste7 = seventh_table.tolist()
        ma_chaine7 = ', '.join(str(element) for element in liste7)
        ma_chaine7 = ma_chaine7.replace("%", "").replace(" ", "")
        sous_chaines7 = ma_chaine7.split(',')
        liste_7 = [int(x) for x in sous_chaines7]
        seventh_tables_data.append(liste_7)
        
        liste8 = eighth_table.tolist()
        ma_chaine8 = ', '.join(str(element) for element in liste8)
        ma_chaine8 = ma_chaine8.replace("%", "").replace(" ", "")
        sous_chaines8 = ma_chaine8.split(',')
        liste_8 = [int(x) for x in sous_chaines8]
        eighth_tables_data.append(liste_8)
        
        liste9 = ninth_table.tolist()
        ma_chaine9 = ', '.join(str(element) for element in liste9)
        ma_chaine9 = ma_chaine9.replace("%", "").replace(" ", "")
        sous_chaines9 = ma_chaine9.split(',')
        liste_9 = [int(x) for x in sous_chaines9]
        ninth_tables_data.append(liste_9)
        
        liste10 = tenth_table.tolist()
        ma_chaine10 = ', '.join(str(element) for element in liste10)
        ma_chaine10 = ma_chaine10.replace("%", "").replace(" ", "")
        sous_chaines10 = ma_chaine10.split(',')
        liste_10 = [int(x) for x in sous_chaines10]
        tenth_tables_data.append(liste_10)
               
        liste15 = fifteenth_table.tolist()
        ma_chaine15 = ', '.join(str(element) for element in liste15)
        ma_chaine15 = ma_chaine15.replace("%", "").replace(" ", "")
        sous_chaines15 = ma_chaine15.split(',')
        liste_15 = [int(x) for x in sous_chaines15]
        fifteenth_tables_data.append(liste_15)
        
        liste16 = sixteenth_table.tolist()
        ma_chaine16 = ', '.join(str(element) for element in liste16)
        ma_chaine16 = ma_chaine16.replace("%", "").replace(" ", "")
        sous_chaines16 = ma_chaine16.split(',')
        liste_16 = [int(x) for x in sous_chaines16]
        sixteenth_tables_data.append(liste_16)

        
column_averages7 = []
for column in zip(*seventh_tables_data):
    column_average7 = mean(column)
    column_averages7.append(column_average7)

column_averages8 = []
for column in zip(*eighth_tables_data):
    column_average8 = mean(column)
    column_averages8.append(column_average8)
    
column_averages9 = []
for column in zip(*ninth_tables_data):
    column_average9 = mean(column)
    column_averages9.append(column_average9)
    
column_averages10 = []
for column in zip(*tenth_tables_data):
    column_average10 = mean(column)
    column_averages10.append(column_average10)
    
column_averages15 = []
for column in zip(*fifteenth_tables_data):
    column_average15 = mean(column)
    column_averages15.append(column_average15)
    
column_averages16 = []
for column in zip(*sixteenth_tables_data):
    column_average16 = mean(column)
    column_averages16.append(column_average16)
    
# enregistrement des données
# Calcul du left_number
left_avg1 = round(column_averages7[2], 1) + round(column_averages7[3], 1) + round(column_averages7[4], 1)
left_avg2 = round(column_averages9[2], 1) + round(column_averages9[3], 1) + round(column_averages9[4], 1)
left_number = max(left_avg1, left_avg2)

# Calcul du right_number
right_avg1 = round(column_averages8[2], 1) + round(column_averages8[3], 1) + round(column_averages8[4], 1)
right_avg2 = round(column_averages8[2], 1) + round(column_averages10[4], 1) + round(column_averages10[3], 1)
right_number = max(right_avg1, right_avg2)

# Calcul du rula_number
rula_avg1 = mean([round(column_averages15[1], 0), round(column_averages15[2], 0)])
rula_avg2 = mean([round(column_averages16[1], 0), round(column_averages16[2], 0)])
rula_number = rula_avg1 + rula_avg2

# Diagramme circulaire
# Données du diagramme circulaire
labels1 = ['Inexistant  <0°', 'Faible  [0° - 10°[', 'Modéré  [10° - 20°[', 'Important  [20° - 30°[', 'Très important  >30°']
labels2 = ['Inexistant  [0° - 10°[', 'Faible  [10° - 20°[', 'Modéré  [20° - 30°[', 'Important  [30° - 40°[', 'Très important  >40°']
labels3 = ['Inexistant  <0°', 'Faible  [0° - 30°[', 'Modéré  [30° - 60°[', 'Important  [60° - 90°[', 'Très important  >90°']
labels4 = ['Inexistant  [0° - 25°[', 'Faible  [25° - 50°[', 'Modéré  [50° - 75°[', 'Important  [75° - 100°[', 'Très important  >100°']
labels5 = ['Inexistant  [0° - 40°[', 'Faible  [40° - 80°[', 'Modéré  [80° - 120°[', 'Important  [120° - 160°[', 'Très important  >160°']
labels5 = ['Inexistant  [0° - 40°[', 'Faible  [40° - 80°[', 'Modéré  [80° - 120°[', 'Important  [120° - 160°[', 'Très important  >160°']
labels6 = ['Score 1-2', 'Score 3-4', 'Score 5-6']

mycolors = ["#8CCB88", "#4CAF50", "#FFC107", "#FF7043",'#C62828']
mycolors2= ["#8CCB88", "#FFC107", "#C62828"]
explode = (0, 0, 0, 0.2, 0.2)
explode2 = (0, 0, 0)

# Créer le diagramme circulaire


fig, ax = plt.subplots()
ax.pie(column_averages7, autopct='%1.1f%%', colors = mycolors, explode=explode)
ax.set_title('Abduction épaule gauche', fontproperties=font, fontsize=20)
# Ajout de la légende et application de la police Cambria
legend = plt.legend(labels4, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)
plt.savefig('abduction_EG.png', format='png', dpi=300, bbox_inches='tight')

fig, ax = plt.subplots()
ax.pie(column_averages8, autopct='%1.1f%%', colors = mycolors, explode=explode)
ax.set_title('Abduction épaule droite', fontproperties=font, fontsize=20)
# Ajout de la légende et application de la police Cambria
legend = plt.legend(labels4, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)
plt.savefig('abduction_ED.png', format='png', dpi=300, bbox_inches='tight')

fig, ax = plt.subplots()
ax.pie(column_averages9, autopct='%1.1f%%', colors = mycolors, explode=explode)
ax.set_title('Flexion épaule gauche', fontproperties=font, fontsize=20)
# Ajout de la légende et application de la police Cambria
legend = plt.legend(labels5, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)
plt.savefig('flexion_EG.png', format='png', dpi=300, bbox_inches='tight')

fig, ax = plt.subplots()
ax.pie(column_averages10, autopct='%1.1f%%', colors = mycolors, explode=explode)
ax.set_title('Flexion épaule droite', fontproperties=font, fontsize=20)
# Ajout de la légende et application de la police Cambria
legend = plt.legend(labels5, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)
plt.savefig('flexion_ED.png', format='png', dpi=300, bbox_inches='tight')


# Création des données pour l'histogramme
scores = ['Score RULA 1-2', 'Score RULA 3-4', 'Score RULA 5-6']
y_pos = np.arange(len(scores))
  
fig, ax = plt.subplots()

# Affichage de l'histogramme
ax.bar(y_pos, column_averages15[:3], align='center', color=mycolors2)
ax.set_xticks(y_pos)
ax.set_xticklabels(scores, fontproperties=font)
ax.set_title('Score RULA épaule gauche', fontproperties=font, fontsize=20)

# Ajout des étiquettes de valeur pour chaque barre
for i, v in enumerate(column_averages15[:3]):
    ax.text(i, v+0.5, str(round(v,1))+"%", ha='center', fontproperties=font)
plt.savefig('RULA_G.png', format='png', dpi=300, bbox_inches='tight')
    
fig, ax = plt.subplots()

# Affichage de l'histogramme
ax.bar(y_pos, column_averages16[:3], align='center', color=mycolors2)
ax.set_xticks(y_pos)
ax.set_xticklabels(scores, fontproperties=font)
ax.set_title('Score RULA des épaule droite', fontproperties=font, fontsize=20)

# Ajout des étiquettes de valeur pour chaque barre
for i, v in enumerate(column_averages16[:3]):
    ax.text(i, v+0.5, str(round(v,1))+"%", ha='center', fontproperties=font)
plt.savefig('RULA_D.png', format='png', dpi=300, bbox_inches='tight')

# Affichage de l'histogramme
#plt.show()

# Enregistrer l'histogramme du score RULA pour le tronc sous format JPEG de haute qualité

def generer_document():
    # Créer un nouveau document Word
    document = Document("nouveau_document.docx")

    # Récupérer les données saisies par l'utilisateur
    entreprise = entreprise_entry.get()
    conseiller = conseiller_entry.get()
    date = date_entry.get()
    beneficiaire = beneficiaire_entry.get()
    duree_sur_place = duree_sur_place_entry.get()
    lieu = lieu_entry.get()
    description_tache = description_tache_entry.get()
    repetition_charge = repetition_charge_entry.get()
    description_posture = description_posture_entry.get()
    equipements = equipements_entry.get()
    environnement = environnement_entry.get()
    temperature = temperature_entry.get()
    type_sol = type_sol_entry.get()
    liste_EPI = liste_EPI_entry.get()
    individuel_equipe = individuel_equipe_entry.get()
    pauses = pauses_entry.get()
    post = post_entry.get()

    # Dictionnaire des balises et leurs valeurs correspondantes
    balises = {
        "##entreprise##": entreprise,
        "##conseiller##": conseiller,
        "##date##": date,
        "##beneficiaire##": beneficiaire,
        "##duree_sur_place##": duree_sur_place,
        "##lieu##": lieu,
        "##description_tache##": description_tache,
        "##repetition_charge##": repetition_charge,
        "##description_posture##": description_posture,
        "##equipements##": equipements,
        "##environnement##": environnement,
        "##temperature##": temperature,
        "##type_sol##": type_sol,
        "##liste_EPI##": liste_EPI,
        "##individuel_equipe##": individuel_equipe,
        "##pauses##": pauses,
        "##poste##": post,
        "##pourcentage_AFG##": str(left_number),
        "##pourcentage_AFD##": str(right_number),
        "##rula_pourcentage##": str(rula_number),
    }

    # Remplacer les balises par les valeurs correspondantes
    remplacer_balises(document, balises)
    
    # Dictionnaire des balises graphiques et leurs fichiers correspondants
    graphiques = {
        "##graphique_abduction_G##": "abduction_EG.png",
        "##abduction_D##": "abduction_ED.png",
        "##graphique_flexion_G##": "flexion_EG.png",
        "##flexion_D##": "flexion_ED.png",
        "##RULA_G##": "RULA_G.png",
        "##RULA_D##": "RULA_D.png",
        "##questionnaire1##": "subplot1.png",
        "##questionnaire2##": "subplot2.png"
    }

    # Ajouter les graphiques à la place des balises
    for paragraph in document.paragraphs:
        for graphique_balise, fichier_image in graphiques.items():
            if graphique_balise in paragraph.text:
                run = paragraph.runs[0]
                run.text = run.text.replace(graphique_balise, "")
                run.add_picture(fichier_image, width=Inches(6))
                break

    # Sauvegarder le nouveau document Word
    document.save("nouveau_document.docx")
    messagebox.showinfo("Succès", "Le nouveau document a été généré avec succès.")

# Créer une fenêtre
root = tk.Tk()
root.title("Générateur de document Word")
root.geometry("600x900")
root.configure(bg="#2C3E50")

# Style des étiquettes et des champs de saisie
label_font = ("Helvetica", 12)
entry_font = ("Helvetica", 12)
label_fg = "#ECF0F1"
entry_bg = "#34495E"
entry_fg = "#ECF0F1"
button_font = ("Helvetica", 14, "bold")
button_bg = "#3498DB"
button_fg = "#ECF0F1"

# Titre de l'application
title_label = tk.Label(root, text="Générateur de document Word", font=("Helvetica", 18, "bold"), fg=label_fg, bg="#2C3E50")
title_label.pack(pady=20)

# Créer un cadre pour organiser les widgets
frame = tk.Frame(root, bg="#2C3E50")
frame.pack(pady=20)

# Fonction pour créer des étiquettes et des champs de saisie
def create_labeled_entry(frame, text, row):
    label = tk.Label(frame, text=text, font=label_font, fg=label_fg, bg="#2C3E50")
    label.grid(row=row, column=0, pady=5, sticky="e")
    entry = tk.Entry(frame, font=entry_font, bg=entry_bg, fg=entry_fg, insertbackground=entry_fg)
    entry.grid(row=row, column=1, pady=5, padx=10)
    return entry

# Créer des étiquettes et des champs de saisie pour chaque balise
entreprise_entry = create_labeled_entry(frame, "Entreprise :", 0)
conseiller_entry = create_labeled_entry(frame, "Conseiller :", 1)
date_entry = create_labeled_entry(frame, "Date :", 2)
beneficiaire_entry = create_labeled_entry(frame, "Bénéficiaire :", 3)
duree_sur_place_entry = create_labeled_entry(frame, "Durée sur place :", 4)
lieu_entry = create_labeled_entry(frame, "Lieu :", 5)
description_tache_entry = create_labeled_entry(frame, "Description de la tâche :", 6)
repetition_charge_entry = create_labeled_entry(frame, "Répétition de charge :", 7)
description_posture_entry = create_labeled_entry(frame, "Description de posture :", 8)
equipements_entry = create_labeled_entry(frame, "Equipements :", 9)
environnement_entry = create_labeled_entry(frame, "Environnement :", 10)
temperature_entry = create_labeled_entry(frame, "Température :", 11)
type_sol_entry = create_labeled_entry(frame, "Type de sol :", 12)
liste_EPI_entry = create_labeled_entry(frame, "Liste des EPI :", 13)
individuel_equipe_entry = create_labeled_entry(frame, "Individuel ou en équipe :", 14)
pauses_entry = create_labeled_entry(frame, "Possibilités de pauses :", 15)
post_entry = create_labeled_entry(frame, "Poste :", 16)

# Bouton de génération du document
generer_button = tk.Button(root, text="Générer Document", font=button_font, bg=button_bg, fg=button_fg, command=generer_document)
generer_button.pack(pady=20)

# Lancer l'application
root.mainloop()