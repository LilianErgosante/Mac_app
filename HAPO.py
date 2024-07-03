import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import docx
from docx.shared import Inches
import os
import docx
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from docx import Document
from tabula import read_pdf
from tabulate import tabulate
from matplotlib import pyplot as plt
from statistics import mean
import numpy as np
import matplotlib.font_manager as fm


###################################################################
# CODE POUR TRAITER LES QUESTIONNAIRES SUBJECTIFS
###################################################################

# Obtenir le chemin d'accès absolu du répertoire courant
current_directory = os.path.dirname(os.path.abspath(__file__))

# Construire le chemin d'accès complet vers le fichier de police
font_path = os.path.join(current_directory, "Cambria-Font-For-Windows.ttf")

# Définir les propriétés de la police
font_properties = FontProperties(fname=font_path)

# Fonction pour calculer le pourcentage des réponses
def calculate_percentage(question):
    total_responses = len(question)
    response_counts = question.value_counts()
    response_percentages = (response_counts / total_responses) * 100
    return response_percentages

# Définir les couleurs pour chaque réponse
colors = {
    "Tout à fait d'accord": '#549e39',
    "Plutôt d'accord": '#8ab833',
    "Moyennement d'accord": '#c0cf3a',
    "Plutôt pas d'accord": '#029676',
    "Pas du tout d'accord": '#4ab5c4'
}

# Charger le fichier Excel et lire les données du questionnaire dans un DataFrame
df1 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="F:G", header=None)
df2 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="H:J", header=None)
df3 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="K:M", header=None)
df4 = pd.read_excel('Questionnaire suivi exosquelette TEST PRESTA AUTO(1-2).xlsx', usecols="N:Q", header=None)

# Supprimer la première ligne (questions) et la remplacer par les colonnes appropriées
df1 = df1.drop(0)
df1.columns = ['Question1', 'Question2']
df2 = df2.drop(0)
df2.columns = ['Question3', 'Question4', 'Question5']
df3 = df3.drop(0)
df3.columns = ['Question6', 'Question7', 'Question8'] 
df4 = df4.drop(0)
df4.columns = ['Question9', 'Question10', 'Question11', 'Question12'] 

# Renommer les réponses contenant des caractères spéciaux
df1.replace({"Plutôt d’accord": "Plutôt d'accord",
             "Pas du tout d’accord": "Pas du tout d'accord",
             "Moyennement d’accord": "Moyennement d'accord",
             "Tout à fait d'accord": "Tout à fait d'accord",
             "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)

# Calculer les pourcentages de réponses pour chaque question
question1_percentages = calculate_percentage(df1['Question1'])
question2_percentages = calculate_percentage(df1['Question2'])
question_percentages_2 = {}
for col in df2.columns:
    df2[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df2.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_2[col] = calculate_percentage(df2[col])

question_percentages_3 = {} # Ajouter cette ligne
for col in df3.columns: # Ajouter cette ligne
    df3[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df3.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_3[col] = calculate_percentage(df3[col]) # Ajouter cette ligne

question_percentages_4 = {} # Ajouter cette ligne
for col in df4.columns: # Ajouter cette ligne
    df4[col].replace({"Plutôt d’accord": "Plutôt d'accord",
                      "Pas du tout d’accord": "Pas du tout d'accord",
                      "Moyennement d’accord": "Moyennement d'accord",
                      "Tout à fait d'accord": "Tout à fait d'accord",
                      "Plutôt pas d’accord": "Plutôt pas d'accord"}, inplace=True)
    df4.replace({"Tout à fait d’accord": "Tout à fait d'accord"}, inplace=True)
    question_percentages_4[col] = calculate_percentage(df4[col]) # Ajouter cette ligne
    
# Créer les subplots
cm = 1/2.54  # centimeters in inches
fig1, (ax1, ax2) = plt.subplots(2, 1, figsize=(20*cm, 30*cm))
fig2, (ax3, ax4) = plt.subplots(2, 1, figsize=(20*cm, 30*cm))

font = FontProperties(fname=font_path)

# Gérer le premier subplot
handles = []
labels = []
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

# Première question
bottom_values_q1 = [0] * len(question1_percentages)
for i, (category, percentage) in enumerate(question1_percentages.items()):
    ax1.bar('Question 1', percentage, bottom=bottom_values_q1, color=colors.get(category, 'lightgray'))
    bottom_values_q1 = [sum(x) for x in zip(bottom_values_q1, [percentage])]

# Deuxième question
bottom_values_q2 = [0] * len(question2_percentages)
for i, (category, percentage) in enumerate(question2_percentages.items()):
    ax1.bar('Question 2', percentage, bottom=bottom_values_q2, color=colors.get(category, 'lightgray'))
    bottom_values_q2 = [sum(x) for x in zip(bottom_values_q2, [percentage])]

#ax1.set_xlabel('Questions', fontproperties=font)
ax1.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax1.set_title('Mise en place et réglages de l\'exosquelette', fontproperties=title_font)

ax1.legend(handles, labels, loc='upper left', bbox_to_anchor=(1.0, 1.0), prop=font)
ax1.grid(True)

ax1.set_xticks([0, 1])
ax1.set_xticklabels(['Facilité de mise en place', 'Adaptation à la morphologie'], fontproperties=font)

# Gérer le deuxième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q3_to_q7 = [0] * len(df2.columns)
for col, percentages in enumerate(question_percentages_2.values()):
    bottom_values_q3_to_q7 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):
        ax2.bar('Question {}'.format(col+3), percentage, bottom=bottom_values_q3_to_q7, color=colors.get(category, 'lightgray'))
        bottom_values_q3_to_q7 = [sum(x) for x in zip(bottom_values_q3_to_q7, [percentage])]

#ax2.set_xlabel('Questions', fontproperties=font)
ax2.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax2.set_title("L'utilisation de l'exosquelette", fontproperties=title_font)

ax2.grid(True)

ax2.set_xticks(range(len(df2.columns)))
ax2.set_xticklabels(['Facilité des mouvements', 'Facilité de déplacement', 'Habituation'], fontproperties=font)

# Enregistrer et afficher la première figure
fig1.tight_layout()
fig1.savefig('subplot1.png')
#plt.show()

# Gérer le troisième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q8_to_q10 = [0] * len(df3.columns)
for col, percentages in enumerate(question_percentages_3.values()):
    bottom_values_q8_to_q10 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):

        ax3.bar('Question {}'.format(col+8), percentage, bottom=bottom_values_q8_to_q10, color=colors.get(category, 'lightgray'))
        bottom_values_q8_to_q10 = [sum(x) for x in zip(bottom_values_q8_to_q10, [percentage])]

#ax3.set_xlabel('Questions', fontproperties=font)
ax3.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax3.set_title("Le confort de l'exosquelette", fontproperties=title_font)

ax3.grid(True)

ax3.set_xticks(range(len(df3.columns)))
ax3.set_xticklabels(['Confort', 'Frottements', 'Me donne chaud'], fontproperties=font)

# Gérer le quatrième subplot
for category, color in colors.items():
    handles.append(plt.Rectangle((0, 0), 1, 1, fc=color))
    labels.append(category)

bottom_values_q11_to_q14 = [0] * len(df4.columns)
for col, percentages in enumerate(question_percentages_4.values()):
    bottom_values_q11_to_q14 = [0] * len(percentages)
    for i, (category, percentage) in enumerate(percentages.items()):
        ax4.bar('Question {}'.format(col+11), percentage, bottom=bottom_values_q11_to_q14, color=colors.get(category, 'lightgray'))
        bottom_values_q11_to_q14 = [sum(x) for x in zip(bottom_values_q11_to_q14, [percentage])]

#ax4.set_xlabel('Questions', fontproperties=font)
ax4.set_ylabel('Pourcentage', fontproperties=font)
title_font = FontProperties(fname=font_path, size=18, weight='bold')
ax4.set_title("Le ressenti avec l'exosquelette", fontproperties=title_font)

ax4.grid(True)

ax4.set_xticks(range(len(df4.columns)))
ax4.set_xticklabels(['Je me sens nerveux', 'J\'aime travailler avec', 'Je me sens isolé avec','Je me sens confiant avec'], fontproperties=font)

# Enregistrer et afficher la deuxième figure
fig2.tight_layout()
fig2.savefig('subplot2.png')
#plt.show()

###################################################################
# CODE POUR IMPORTER LA PHOTO
###################################################################

def replace_photo_tag_with_image(doc_path, image_path, new_doc_path):
    # Ouvrir le document Word
    doc = docx.Document(doc_path)

    # Parcourir les paragraphes du document
    for paragraph in doc.paragraphs:
        # Vérifier si le paragraphe contient la balise ##photo##
        if "##photo##" in paragraph.text:
            # Remplacer la balise par une image
            run = paragraph.runs[0]
            run.text = ""  # Effacer la balise
            run.add_picture(image_path, width=Inches(5))  # Ajouter l'image

    # Enregistrer les modifications dans le nouveau document
    doc.save(new_doc_path)

# Créer une fenêtre tkinter cachée pour ouvrir la boîte de dialogue
root = tk.Tk()
root.withdraw()
root.deiconify()  # Forcer l'affichage de la fenêtre

# Ouvrir la boîte de dialogue pour sélectionner l'image
image_path = filedialog.askopenfilename(title="Sélectionner une image", filetypes=[("Fichiers image", "*.png *.jpg *.jpeg")])

# Vérifier si une image a été sélectionnée
if image_path:
    # Appeler la fonction pour remplacer la balise ##photo## par l'image dans un nouveau document
    new_doc_path = "nouveau_document.docx"
    replace_photo_tag_with_image("CR_HAPO.docx", image_path, new_doc_path)
    print(f"L'image a été insérée avec succès dans le nouveau document : {new_doc_path}")
else:
    print("Aucune image sélectionnée.")

# Fermer la fenêtre tkinter cachée
root.destroy()


###################################################################
# REMPLISSAGE DU RAPPORT
###################################################################

def remplacer_balises(document, balises):
    """
    Remplace les balises dans le document par les valeurs correspondantes.

    Arguments :
    document -- Document Word (obj)
    balises -- Dictionnaire contenant les balises en tant que clés et leurs valeurs correspondantes en tant que valeurs (dict)
    """
    for paragraph in document.paragraphs:
        for balise, valeur in balises.items():
            if balise in paragraph.text:
                paragraph.text = paragraph.text.replace(balise, valeur)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for balise, valeur in balises.items():
                        if balise in paragraph.text:
                            paragraph.text = paragraph.text.replace(balise, valeur)
                            
                            
# Ouvrir les boîtes de dialogue pour sélectionner les dossiers
directory_sans = filedialog.askdirectory(title="Sélectionner un dossier pour les rapports sans exosquelette")
directory_avec = filedialog.askdirectory(title="Sélectionner un dossier pour les rapports avec exosquelette")

# Vérifier si les deux dossiers ont été sélectionnés
if directory_sans and directory_avec:
    # Obtenez le chemin absolu des dossiers sélectionnés
    pdf_folder_sans = os.path.abspath(directory_sans)
    pdf_folder_avec = os.path.abspath(directory_avec)
    print(f"Chemin absolu du dossier sélectionné pour sans exosquelette : {pdf_folder_sans}")
    print(f"Chemin absolu du dossier sélectionné pour avec exosquelette : {pdf_folder_avec}")
else:
    print("Un ou deux dossiers n'ont pas été sélectionnés.")
    # Gérer le cas où un ou les deux dossiers n'ont pas été sélectionnés

fourth_tables_data = []
fourteenth_tables_data = []

for file_name in os.listdir(pdf_folder_sans):
    if file_name.endswith('.pdf'):
        file_path = os.path.join(pdf_folder_sans, file_name)
        dfs = read_pdf(file_path, pages='all', multiple_tables=True)
        fourth_table = dfs[3].iloc[:, 1]
        fourteenth_table = dfs[13].iloc[:, 1]

        liste4 = fourth_table.tolist()
        ma_chaine4 = ', '.join(str(element) for element in liste4)
        ma_chaine4 = ma_chaine4.replace("%", "").replace(" ", "")
        sous_chaines4 = ma_chaine4.split(',')
        liste_4 = [int(x) for x in sous_chaines4]
        fourth_tables_data.append(liste_4)
        
        liste14 = fourteenth_table.tolist()
        ma_chaine14 = ', '.join(str(element) for element in liste14)
        ma_chaine14 = ma_chaine14.replace("%", "").replace(" ", "")
        sous_chaines14 = ma_chaine14.split(',')
        liste_14 = [int(x) for x in sous_chaines14]
        fourteenth_tables_data.append(liste_14)

column_averages4_sans = []
for column in zip(*fourth_tables_data):
    column_average4_sans = mean(column)
    column_averages4_sans.append(column_average4_sans)
    
column_averages14_sans = []
for column in zip(*fourteenth_tables_data):
    column_average14_sans = mean(column)
    column_averages14_sans.append(column_average14_sans)
    
# enregistrement des données
second_number_sans = round(column_averages4_sans[1], 1)
third_number_sans = round(column_averages4_sans[2], 0)+round(column_averages4_sans[3], 1)+round(column_averages4_sans[4], 0)
rula_back_sans = round(column_averages14_sans[1], 0) + round(column_averages14_sans[2], 0)

# Diagramme circulaire
labels3 = ['Inexistant  <0°', 'Faible  [0° - 30°[', 'Modéré  [30° - 60°[', 'Important  [60° - 90°[', 'Très important  >90°']
mycolors = ["#8CCB88", "#4CAF50", "#FFC107", "#FF7043",'#C62828']
mycolors2= ["#8CCB88", "#FFC107", "#C62828"]
explode = (0, 0, 0, 0.2, 0.2)
font = FontProperties(fname=font_path)

fig, ax = plt.subplots()
ax.pie(column_averages4_sans, autopct='%1.1f%%', colors=mycolors, explode=explode)
ax.set_title('Flexion de tronc sans l\'exosquelette', fontproperties=font, fontsize=20)
legend = plt.legend(labels3, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)

plt.savefig("graphique_sans.png", dpi=300, bbox_inches='tight')

# Création des données pour l'histogramme
scores = ['Score RULA 1-2', 'Score RULA 3-4', 'Score RULA 5-6']
y_pos = np.arange(len(scores))

# Création de la figure et des axes
fig, ax = plt.subplots(figsize=(8, 6))

# Affichage de l'histogramme pour le score RULA du tronc
ax.bar(y_pos, column_averages14_sans[:3], align='center', color=mycolors2)
ax.set_xticks(y_pos)
ax.set_xticklabels(scores, fontproperties=font)
ax.set_title('Score RULA du tronc sans l\'exosquelette', fontproperties=font, fontsize=20)

# Ajout des étiquettes de valeur pour chaque barre
for i, v in enumerate(column_averages14_sans[:3]):
    ax.text(i, v+0.5, str(round(v,1))+"%", ha='center', fontproperties=font)

# Ajustement des marges pour s'assurer que tout le contenu est visible
plt.tight_layout()

# Enregistrer l'histogramme du score RULA pour le tronc sous format JPEG de haute qualité
plt.savefig('rula_sans.png', format='png', dpi=300)

fourth_tables_data2 = []
fourteenth_tables_data2 = []

for file_name in os.listdir(pdf_folder_avec):
    if file_name.endswith('.pdf'):
        file_path = os.path.join(pdf_folder_avec, file_name)
        dfs = read_pdf(file_path, pages='all', multiple_tables=True)
        fourth_table2 = dfs[3].iloc[:, 1]
        fourteenth_table2 = dfs[13].iloc[:, 1]

        liste4 = fourth_table2.tolist()
        ma_chaine4 = ', '.join(str(element) for element in liste4)
        ma_chaine4 = ma_chaine4.replace("%", "").replace(" ", "")
        sous_chaines4 = ma_chaine4.split(',')
        liste_4 = [int(x) for x in sous_chaines4]
        fourth_tables_data2.append(liste_4)
        
        liste14 = fourteenth_table2.tolist()
        ma_chaine14 = ', '.join(str(element) for element in liste14)
        ma_chaine14 = ma_chaine14.replace("%", "").replace(" ", "")
        sous_chaines14 = ma_chaine14.split(',')
        liste_14 = [int(x) for x in sous_chaines14]
        fourteenth_tables_data2.append(liste_14)

column_averages4_avec = []
for column in zip(*fourth_tables_data2):
    column_average4_avec = mean(column)
    column_averages4_avec.append(column_average4_avec)
    
column_averages14_avec = []
for column in zip(*fourteenth_tables_data2):
    column_average14_avec = mean(column)
    column_averages14_avec.append(column_average14_avec)
    
# enregistrement des données
second_number_avec = round(column_averages4_avec[1], 1)
third_number_avec = round(column_averages4_avec[2], 0)+round(column_averages4_avec[3], 1)+round(column_averages4_avec[4], 0)
rula_back_avec = round(column_averages14_avec[1], 0) + round(column_averages14_avec[2], 0)
diff1 = third_number_sans - third_number_avec
diff_rula = rula_back_sans - rula_back_avec
# Diagramme circulaire

fig, ax = plt.subplots()
ax.pie(column_averages4_avec, autopct='%1.1f%%', colors=mycolors, explode=explode)
ax.set_title('Flexion de tronc avec l\'exosquelette', fontproperties=font, fontsize=20)
legend = plt.legend(labels3, title='Légende', loc='center left', bbox_to_anchor=(1, 0.5), prop=font)
legend.set_title('Légende', prop=font)

plt.savefig("graphique_avec.png", dpi=300, bbox_inches='tight')

# Création des données pour l'histogramme
scores = ['Score RULA 1-2', 'Score RULA 3-4', 'Score RULA 5-6']
y_pos = np.arange(len(scores))

# Création de la figure et des axes
fig, ax = plt.subplots(figsize=(8, 6))

# Affichage de l'histogramme pour le score RULA du tronc
ax.bar(y_pos, column_averages14_avec[:3], align='center', color=mycolors2)
ax.set_xticks(y_pos)
ax.set_xticklabels(scores, fontproperties=font)
ax.set_title('Score RULA du tronc avec l\'exosquelette', fontproperties=font, fontsize=20)

# Ajout des étiquettes de valeur pour chaque barre
for i, v in enumerate(column_averages14_avec[:3]):
    ax.text(i, v+0.5, str(round(v,1))+"%", ha='center', fontproperties=font)

# Ajustement des marges pour s'assurer que tout le contenu est visible
plt.tight_layout()

# Enregistrer l'histogramme du score RULA pour le tronc sous format JPEG de haute qualité
plt.savefig('rula_avec.png', format='png', dpi=300)

def generer_document():
    document = Document("nouveau_document.docx")

    entreprise = entreprise_entry.get()
    conseiller = conseiller_entry.get()
    date = date_entry.get()
    beneficiaire = beneficiaire_entry.get()
    duree_sur_place = duree_sur_place_entry.get()
    lieu = lieu_entry.get()
    description_tache = description_tache_entry.get()
    repetition_charge = repetition_charge_entry.get()
    description_posture = description_posture_entry.get()
    equipements = equipements_entry.get()
    environnement = environnement_entry.get()
    temperature = temperature_entry.get()
    type_sol = type_sol_entry.get()
    liste_EPI = liste_EPI_entry.get()
    individuel_equipe = individuel_equipe_entry.get()
    pauses = pauses_entry.get()
    post = post_entry.get()

    balises = {
        "##entreprise##": entreprise,
        "##conseiller##": conseiller,
        "##date##": date,
        "##beneficiaire##": beneficiaire,
        "##duree_sur_place##": duree_sur_place,
        "##lieu##": lieu,
        "##description_tache##": description_tache,
        "##repetition_charge##": repetition_charge,
        "##description_posture##": description_posture,
        "##equipements##": equipements,
        "##environnement##": environnement,
        "##temperature##": temperature,
        "##type_sol##": type_sol,
        "##liste_EPI##": liste_EPI,
        "##individuel_equipe##": individuel_equipe,
        "##pauses##": pauses,
        "##poste##": post,
        "##pourcentage##": str(second_number_sans),
        "##rula_pourcentage##": str(rula_back_sans),
        "##diff##": str(diff1),
        "##diff_rula##": str(diff_rula)
        
    }

    remplacer_balises(document, balises)
    
    graphique_balise = "##graphique_flexion_tronc_sans##"
    graphique_balise2 = "##rula_tronc_sans##"
    graphique_balise3 = "##questionnaire1##"
    graphique_balise4 = "##questionnaire2##"
    graphique_balise5 = "##graphique_flexion_tronc_avec##"
    graphique_balise6 = "##rula_tronc_avec##"

    for paragraph in document.paragraphs:
        if graphique_balise in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise, "")
            run.add_picture("graphique_sans.png", width=Inches(6))
            break
            
    for paragraph in document.paragraphs:
        if graphique_balise2 in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise2, "")
            run.add_picture("rula_sans.png", width=Inches(6))
            break
            
    for paragraph in document.paragraphs:
        if graphique_balise3 in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise3, "")
            run.add_picture("subplot1.png", width=Inches(6))
            break
            
    for paragraph in document.paragraphs:
        if graphique_balise4 in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise4, "")
            run.add_picture("subplot2.png", width=Inches(6))
            break
        
    for paragraph in document.paragraphs:
        if graphique_balise5 in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise5, "")
            run.add_picture("graphique_avec.png", width=Inches(6))
            break
        
    for paragraph in document.paragraphs:
        if graphique_balise6 in paragraph.text:
            run = paragraph.runs[0]
            run.text = run.text.replace(graphique_balise6, "")
            run.add_picture("rula_avec.png", width=Inches(6))
            break

    document.save("nouveau_document.docx")
    messagebox.showinfo("Succès", "Le nouveau document a été généré avec succès.")

# Créer une fenêtre
root = tk.Tk()
root.title("Générateur de document Word")
root.geometry("600x900")
root.configure(bg="#2C3E50")

# Style des étiquettes et des champs de saisie
label_font = ("Helvetica", 12)
entry_font = ("Helvetica", 12)
label_fg = "#ECF0F1"
entry_bg = "#34495E"
entry_fg = "#ECF0F1"
button_font = ("Helvetica", 14, "bold")
button_bg = "#3498DB"
button_fg = "#ECF0F1"

# Titre de l'application
title_label = tk.Label(root, text="Générateur de document Word", font=("Helvetica", 18, "bold"), fg=label_fg, bg="#2C3E50")
title_label.pack(pady=20)

# Créer un cadre pour organiser les widgets
frame = tk.Frame(root, bg="#2C3E50")
frame.pack(pady=20)

# Fonction pour créer des étiquettes et des champs de saisie
def create_labeled_entry(frame, text, row):
    label = tk.Label(frame, text=text, font=label_font, fg=label_fg, bg="#2C3E50")
    label.grid(row=row, column=0, pady=5, sticky="e")
    entry = tk.Entry(frame, font=entry_font, bg=entry_bg, fg=entry_fg, insertbackground=entry_fg)
    entry.grid(row=row, column=1, pady=5, padx=10)
    return entry

# Créer des étiquettes et des champs de saisie pour chaque balise
entreprise_entry = create_labeled_entry(frame, "Entreprise :", 0)
conseiller_entry = create_labeled_entry(frame, "Conseiller :", 1)
date_entry = create_labeled_entry(frame, "Date :", 2)
beneficiaire_entry = create_labeled_entry(frame, "Bénéficiaire :", 3)
duree_sur_place_entry = create_labeled_entry(frame, "Durée sur place :", 4)
lieu_entry = create_labeled_entry(frame, "Lieu :", 5)
description_tache_entry = create_labeled_entry(frame, "Description de la tâche :", 6)
repetition_charge_entry = create_labeled_entry(frame, "Répétition de charge :", 7)
description_posture_entry = create_labeled_entry(frame, "Description de posture :", 8)
equipements_entry = create_labeled_entry(frame, "Equipements :", 9)
environnement_entry = create_labeled_entry(frame, "Environnement :", 10)
temperature_entry = create_labeled_entry(frame, "Température :", 11)
type_sol_entry = create_labeled_entry(frame, "Type de sol :", 12)
liste_EPI_entry = create_labeled_entry(frame, "Liste des EPI :", 13)
individuel_equipe_entry = create_labeled_entry(frame, "Individuel ou en équipe :", 14)
pauses_entry = create_labeled_entry(frame, "Possibilités de pauses :", 15)
post_entry = create_labeled_entry(frame, "Poste :", 16)

# Bouton de génération du document
generer_button = tk.Button(root, text="Générer Document", font=button_font, bg=button_bg, fg=button_fg, command=generer_document)
generer_button.pack(pady=20)

# Lancer l'application
root.mainloop()
